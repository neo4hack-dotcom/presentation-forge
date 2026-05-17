"""Document parsing for context ingestion."""
from __future__ import annotations

import io
from pathlib import Path

from pypdf import PdfReader
from docx import Document
import markdown as md
from bs4 import BeautifulSoup


MAX_CHARS_PER_DOC = 60_000


def parse_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n\n".join(parts).strip()


def parse_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    out = []
    for p in doc.paragraphs:
        if p.text.strip():
            out.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                out.append(" | ".join(cells))
    return "\n".join(out).strip()


def parse_markdown(data: bytes) -> str:
    text = data.decode("utf-8", errors="replace")
    html = md.markdown(text)
    soup = BeautifulSoup(html, "html.parser")
    return soup.get_text("\n").strip()


def parse_text(data: bytes) -> str:
    return data.decode("utf-8", errors="replace").strip()


def parse_file(name: str, data: bytes) -> str:
    ext = Path(name).suffix.lower()
    try:
        if ext == ".pdf":
            text = parse_pdf(data)
        elif ext in (".docx",):
            text = parse_docx(data)
        elif ext in (".md", ".markdown"):
            text = parse_markdown(data)
        else:
            text = parse_text(data)
    except Exception as e:
        return f"[Could not parse {name}: {e}]"
    if len(text) > MAX_CHARS_PER_DOC:
        text = text[:MAX_CHARS_PER_DOC] + f"\n\n[...truncated, original {len(text)} chars]"
    return text
